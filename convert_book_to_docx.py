#!/usr/bin/env python3
"""
Convert Markdown book to DOCX with KDP formatting:
- Garamond 11pt for main content
- Monospace for code
- TOC generation
- 6x9 KDP format
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    import markdown
    from markdown.extensions import codehilite, fenced_code, tables, toc
except ImportError:
    print("Installing required packages...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "markdown", "pygments"])
    import markdown
    from markdown.extensions import codehilite, fenced_code, tables, toc

def setup_document_styles(doc):
    """Configure document styles for KDP 6x9 format"""
    # Set page size to 6x9 inches (KDP format)
    section = doc.sections[0]
    section.page_height = Inches(9)
    section.page_width = Inches(6)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    
    # Configure Normal style (Garamond 11pt)
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Garamond'
    normal_font.size = Pt(11)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(0)  # No extra space to prevent double breaks
    
    # Configure Heading styles with Garamond
    for i in range(1, 7):
        heading_style = doc.styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = 'Garamond'
        heading_font.size = Pt(14 - i)  # Decreasing sizes for lower headings
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(0)  # No extra space after headings
    
    # Create code style (monospace)
    try:
        code_style = doc.styles['Code']
    except KeyError:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    
    code_font = code_style.font
    code_font.name = 'Courier New'
    code_font.size = Pt(9)
    code_style.paragraph_format.left_indent = Inches(0.25)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(0)  # No extra space after code blocks
    
    # Create inline code style
    try:
        inline_code_style = doc.styles['Inline Code']
    except KeyError:
        inline_code_style = doc.styles.add_style('Inline Code', WD_STYLE_TYPE.CHARACTER)
    
    inline_code_font = inline_code_style.font
    inline_code_font.name = 'Courier New'
    inline_code_font.size = Pt(10)

def add_table_of_contents(doc, headings):
    """Add a table of contents page"""
    doc.add_page_break()
    para = doc.add_paragraph('Table of Contents')
    para.style = 'Heading 1'
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing
    
    for level, text, page_num in headings:
        toc_para = doc.add_paragraph()
        toc_para.style = 'Normal'
        
        # Add indentation based on level
        if level == 1:
            toc_para.paragraph_format.left_indent = Inches(0)
            toc_run = toc_para.add_run(text)
            toc_run.bold = True
            toc_run.font.size = Pt(11)
        elif level == 2:
            toc_para.paragraph_format.left_indent = Inches(0.25)
            toc_run = toc_para.add_run(text)
            toc_run.font.size = Pt(11)
        elif level == 3:
            toc_para.paragraph_format.left_indent = Inches(0.5)
            toc_run = toc_para.add_run(f"  {text}")
            toc_run.font.size = Pt(10)
        else:
            toc_para.paragraph_format.left_indent = Inches(0.75)
            toc_run = toc_para.add_run(f"    {text}")
            toc_run.font.size = Pt(10)
        
        # Add dots and page number (simplified - actual page numbers require field codes)
        dots = '.' * (50 - len(text))
        toc_para.add_run(f" {dots} ").font.size = Pt(9)
        toc_para.add_run(str(page_num)).font.size = Pt(9)
    
    doc.add_page_break()

def parse_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to DOCX with proper formatting"""
    print(f"Reading {md_file}...")
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Check if TOC already exists
    has_toc = re.search(r'^#+\s*Table\s+of\s+Contents', md_content, re.IGNORECASE | re.MULTILINE)
    
    # Extract headings for TOC
    headings = []
    heading_pattern = r'^(#{1,6})\s+(.+)$'
    for match in re.finditer(heading_pattern, md_content, re.MULTILINE):
        level = len(match.group(1))
        text = match.group(2).strip()
        # Skip TOC heading itself
        if 'table of contents' not in text.lower():
            headings.append((level, text, 1))  # Page numbers would need actual rendering
    
    # Create document
    doc = Document()
    setup_document_styles(doc)
    
    # Add TOC if not present
    if not has_toc and headings:
        print("Adding Table of Contents...")
        add_table_of_contents(doc, headings[:30])  # Limit to first 30 for TOC
    
    # Parse markdown
    lines = md_content.split('\n')
    in_code_block = False
    code_block_lines = []
    code_language = ''
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Skip TOC section if it exists (we'll add our own)
        if re.match(r'^#+\s*Table\s+of\s+Contents', line, re.IGNORECASE):
            # Skip until next major heading
            i += 1
            while i < len(lines) and not re.match(r'^#+\s+', lines[i]):
                i += 1
            continue
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                if code_block_lines:
                    code_para = doc.add_paragraph('\n'.join(code_block_lines))
                    code_para.style = 'Code'
                    code_para.paragraph_format.space_after = Pt(0)
                code_block_lines = []
                in_code_block = False
                code_language = ''
            else:
                # Start code block
                in_code_block = True
                code_language = line[3:].strip()
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Handle headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Remove markdown formatting from heading text
            text = re.sub(r'\*\*([^\*]+)\*\*', r'\1', text)  # Remove bold markers, keep text
            text = re.sub(r'\*([^\*]+)\*', r'\1', text)  # Remove italic markers, keep text
            para = doc.add_heading(text, level=min(level, 6))
            para.paragraph_format.space_after = Pt(0)
            i += 1
            continue
        
        # Handle horizontal rules (but not list items that start with *)
        if re.match(r'^\*+\s*$|^-{3,}\s*$|^_{3,}\s*$', line) and not re.match(r'^[\*\-\+]\s+', line):
            # Add a horizontal rule as a paragraph with border
            hr_para = doc.add_paragraph()
            hr_para.paragraph_format.space_before = Pt(6)
            hr_para.paragraph_format.space_after = Pt(6)
            hr_run = hr_para.add_run('_' * 50)
            hr_run.font.size = Pt(8)
            i += 1
            continue
        
        # Handle lists (must have space after bullet/number)
        if re.match(r'^[\*\-\+]\s+', line) or re.match(r'^\d+\.\s+', line):
            # Remove list marker and process content
            list_text = re.sub(r'^[\*\-\+]\s+|^\d+\.\s+', '', line)
            para = doc.add_paragraph(style='List Bullet' if re.match(r'^[\*\-\+]\s+', line) else 'List Number')
            para.paragraph_format.space_after = Pt(0)
            # Process formatting in list item
            process_inline_formatting(para, list_text)
            i += 1
            continue
        
        # Handle regular paragraphs
        if line.strip():
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(0)  # No extra space
            process_inline_formatting(para, line)
        else:
            # Empty line - add minimal spacing
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(0)
        
        i += 1
    
    # Save document
    print(f"Saving to {docx_file}...")
    doc.save(docx_file)
    print(f"Conversion complete! Saved to {docx_file}")

def process_inline_formatting(para, text):
    """Process inline markdown formatting (bold, italic, code)"""
    # Split by inline code first
    parts = re.split(r'(`[^`]+`)', text)
    
    for part in parts:
        if part.startswith('`') and part.endswith('`'):
            # Inline code
            code_text = part[1:-1]
            run = para.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            # Process bold and italic
            process_bold_italic(para, part)

def process_bold_italic(para, text):
    """Process bold (**text**) and italic (*text*) formatting"""
    # Process text character by character, tracking bold and italic state
    # This handles nested and mixed formatting correctly
    
    i = 0
    current_text = ""
    in_bold = False
    in_italic = False
    in_code = False
    
    while i < len(text):
        # Check for inline code first (has priority)
        if text[i] == '`' and i + 1 < len(text) and text[i+1] != '`':
            # Flush current text
            if current_text:
                run = para.add_run(current_text)
                if in_bold:
                    run.bold = True
                if in_italic:
                    run.italic = True
                run.font.name = 'Garamond'
                current_text = ""
            
            # Find end of code
            end_code = text.find('`', i + 1)
            if end_code != -1:
                code_text = text[i+1:end_code]
                run = para.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                i = end_code + 1
                continue
        
        # Check for bold (**text**)
        elif text[i:i+2] == '**' and not in_code:
            # Flush current text before changing state
            if current_text:
                run = para.add_run(current_text)
                if in_bold:
                    run.bold = True
                if in_italic:
                    run.italic = True
                run.font.name = 'Garamond'
                current_text = ""
            
            # Toggle bold state
            in_bold = not in_bold
            i += 2
            continue
        
        # Check for italic (*text*) - but not if it's part of **
        elif text[i] == '*' and not in_code and (i == 0 or text[i-1] != '*') and (i+1 >= len(text) or text[i+1] != '*'):
            # Flush current text before changing state
            if current_text:
                run = para.add_run(current_text)
                if in_bold:
                    run.bold = True
                if in_italic:
                    run.italic = True
                run.font.name = 'Garamond'
                current_text = ""
            
            # Toggle italic state
            in_italic = not in_italic
            i += 1
            continue
        
        else:
            # Regular character - add to current text
            current_text += text[i]
            i += 1
    
    # Flush any remaining text
    if current_text:
        run = para.add_run(current_text)
        if in_bold:
            run.bold = True
        if in_italic:
            run.italic = True
        run.font.name = 'Garamond'

if __name__ == '__main__':
    md_file = Path('Book/Complete_Shadow_AI_Detection_Book_CREDIBLE.md')
    docx_file = Path('Book/Shadow_AI_Detection_KDP.docx')
    
    if not md_file.exists():
        print(f"Error: {md_file} not found!")
        sys.exit(1)
    
    # Try to delete old file if it exists and is locked, save to temp first
    import shutil
    temp_file = Path('Book/Shadow_AI_Detection_KDP_temp.docx')
    if docx_file.exists():
        try:
            docx_file.unlink()
        except PermissionError:
            print(f"Warning: Could not delete {docx_file} (file may be open).")
            print(f"Saving to {temp_file} instead. Please close the original file and rename.")
            docx_file = temp_file
    
    parse_markdown_to_docx(md_file, docx_file)

